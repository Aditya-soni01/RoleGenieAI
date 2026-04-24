"""
Template service — PDF and DOCX builders for all 10 resume templates.

Each template has a unique visual style while remaining ATS-safe:
- Standard text headings (no images/tables for core body)
- Parseable section labels
- Clean bullet structure
- No complex graphics for content sections

Dispatcher:
    build_pdf(template_id, data)  → bytes
    build_docx(template_id, data) → bytes
"""

import io
from typing import Any, Callable, Dict, List, Optional


# ─── Data type ────────────────────────────────────────────────────────────────

NormalizedData = Dict[str, Any]
"""
Expected keys after _normalize_data():
  full_name, contact, summary, skills, experience, education, projects, certifications
"""


# ═══════════════════════════════════════════════════════════════════════════════
# DOCX BUILDERS
# ═══════════════════════════════════════════════════════════════════════════════

def _docx_base(
    data: NormalizedData,
    *,
    name_size: int = 22,
    name_centered: bool = True,
    accent_rgb: tuple = (0x1A, 0x56, 0xDB),
    section_uppercase: bool = True,
    hr_thickness: str = "6",
    section_font_size: int = 11,
    body_font_size: int = 10,
    compact: bool = False,
    skills_first: bool = False,
    projects_second: bool = False,
    show_skills_grid: bool = False,
) -> bytes:
    """
    Shared DOCX builder parameterised for all 10 template variants.
    """
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    r, g, b = accent_rgb

    doc = Document()
    sec = doc.sections[0]
    top_margin = 0.55 if compact else 0.75
    sec.top_margin = Inches(top_margin)
    sec.bottom_margin = Inches(top_margin)
    sec.left_margin = Inches(0.85 if compact else 1.0)
    sec.right_margin = Inches(0.85 if compact else 1.0)
    doc.styles["Normal"].paragraph_format.space_after = Pt(0)

    def _add_hr(paragraph):
        pPr = paragraph._element.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), hr_thickness)
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), f"{r:02X}{g:02X}{b:02X}")
        pBdr.append(bottom)
        pPr.append(pBdr)

    def heading(title: str):
        label = title.upper() if section_uppercase else title
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8 if compact else 10)
        run = p.add_run(label)
        run.bold = True
        run.font.size = Pt(section_font_size)
        run.font.color.rgb = RGBColor(r, g, b)
        _add_hr(p)

    # ── Name ──────────────────────────────────────────────────────────────────
    # Root-cause fix: doc-level space_after=Pt(0) suppresses all inter-paragraph
    # gaps. Override explicitly on the name paragraph so the contact line never
    # collides with the name glyph, regardless of font size or template variant.
    p_name = doc.add_paragraph()
    p_name.alignment = WD_ALIGN_PARAGRAPH.CENTER if name_centered else WD_ALIGN_PARAGRAPH.LEFT
    p_name.paragraph_format.space_after = Pt(3)          # explicit gap after name
    p_name.paragraph_format.space_before = Pt(0)
    rn = p_name.add_run(data.get("full_name", ""))
    rn.bold = True
    rn.font.size = Pt(name_size)

    # ── Contact ───────────────────────────────────────────────────────────────
    p_contact = doc.add_paragraph()
    p_contact.alignment = WD_ALIGN_PARAGRAPH.CENTER if name_centered else WD_ALIGN_PARAGRAPH.LEFT
    p_contact.paragraph_format.space_after = Pt(6)       # breathing room before first section
    p_contact.paragraph_format.space_before = Pt(0)
    rc = p_contact.add_run(data.get("contact", ""))
    rc.font.size = Pt(body_font_size - 1)

    def _add_summary():
        if data.get("summary"):
            heading("Professional Summary")
            p = doc.add_paragraph(data["summary"])
            p.runs[0].font.size = Pt(body_font_size)

    def _add_skills():
        tech = data.get("technical_skills") or []
        prof = data.get("professional_skills") or []
        all_skills = data.get("skills") or tech + prof
        if all_skills:
            heading("Skills")
            if tech and prof:
                p_tech = doc.add_paragraph()
                r_label = p_tech.add_run("Technical: ")
                r_label.bold = True
                r_label.font.size = Pt(body_font_size)
                r_tech = p_tech.add_run(" • ".join(tech))
                r_tech.font.size = Pt(body_font_size)
                p_prof = doc.add_paragraph()
                r_label2 = p_prof.add_run("Professional: ")
                r_label2.bold = True
                r_label2.font.size = Pt(body_font_size)
                r_prof = p_prof.add_run(" • ".join(prof))
                r_prof.font.size = Pt(body_font_size)
            elif show_skills_grid:
                mid = (len(all_skills) + 1) // 2
                for pair in zip(all_skills[:mid], all_skills[mid:] + [""]):
                    row_text = "  •  ".join(s for s in pair if s)
                    p = doc.add_paragraph(row_text)
                    p.runs[0].font.size = Pt(body_font_size)
            else:
                p = doc.add_paragraph(" • ".join(all_skills))
                p.runs[0].font.size = Pt(body_font_size)

    def _add_experience():
        if data.get("experience"):
            heading("Experience")
            for exp in data["experience"]:
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(5 if compact else 6)
                r_run = p.add_run(f"{exp.get('title', '')}  —  {exp.get('company', '')}")
                r_run.bold = True
                r_run.font.size = Pt(body_font_size)
                p2 = doc.add_paragraph()
                r2 = p2.add_run(exp.get("duration", ""))
                r2.italic = True
                r2.font.size = Pt(body_font_size - 1)
                sub_projects = exp.get("projects") or []
                if sub_projects:
                    for proj in sub_projects:
                        proj_name = proj.get("name", "")
                        if proj_name:
                            p_proj = doc.add_paragraph()
                            r_proj = p_proj.add_run(proj_name)
                            r_proj.bold = True
                            r_proj.italic = True
                            r_proj.font.size = Pt(body_font_size - 1)
                        for bullet in proj.get("bullets") or []:
                            bp = doc.add_paragraph(style="List Bullet")
                            bp.add_run(bullet).font.size = Pt(body_font_size)
                else:
                    for ach in exp.get("achievements", []):
                        bp = doc.add_paragraph(style="List Bullet")
                        bp.add_run(ach).font.size = Pt(body_font_size)

    def _add_education():
        if data.get("education"):
            heading("Education")
            for edu in data["education"]:
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(4)
                r_run = p.add_run(edu.get("degree", ""))
                r_run.bold = True
                r_run.font.size = Pt(body_font_size)
                doc.add_paragraph(
                    f"{edu.get('institution', '')} | {edu.get('year', '')}"
                ).runs[0].font.size = Pt(body_font_size)

    def _add_projects():
        if data.get("projects"):
            heading("Projects")
            for proj in data["projects"]:
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(4)
                r_run = p.add_run(proj.get("name", ""))
                r_run.bold = True
                r_run.font.size = Pt(body_font_size)
                desc = proj.get("description", "")
                if desc:
                    doc.add_paragraph(desc).runs[0].font.size = Pt(body_font_size)
                techs = proj.get("technologies", [])
                if techs:
                    p3 = doc.add_paragraph(f"Technologies: {', '.join(techs)}")
                    p3.runs[0].italic = True
                    p3.runs[0].font.size = Pt(body_font_size - 1)

    def _add_certifications():
        if data.get("certifications"):
            heading("Certifications")
            for cert in data["certifications"]:
                bp = doc.add_paragraph(style="List Bullet")
                bp.add_run(cert).font.size = Pt(body_font_size)

    # ── Section ordering ──────────────────────────────────────────────────────
    if skills_first:
        _add_summary()
        _add_skills()
        _add_experience()
        _add_projects()
    elif projects_second:
        _add_summary()
        _add_projects()
        _add_experience()
        _add_skills()
    else:
        _add_summary()
        _add_skills()
        _add_experience()

    _add_education()

    if not skills_first and not projects_second:
        _add_projects()

    _add_certifications()

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ─── 10 DOCX template dispatchers ────────────────────────────────────────────

def _docx_sidebar(data: NormalizedData) -> bytes:
    """Two-column DOCX layout with skills/contact sidebar and main resume body."""
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(0.6)
    sec.bottom_margin = Inches(0.6)
    sec.left_margin = Inches(0.65)
    sec.right_margin = Inches(0.65)

    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(2.0)
    table.columns[1].width = Inches(5.0)
    sidebar = table.cell(0, 0)
    main = table.cell(0, 1)
    sidebar.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    main.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP

    def shade(cell, fill: str):
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), fill)
        tc_pr.append(shd)

    shade(sidebar, "DBEAFE")

    def para(cell, text: str = "", size: int = 10, bold: bool = False,
             color: tuple = (0x11, 0x18, 0x27), before: int = 0,
             after: int = 2):
        p = cell.add_paragraph()
        p.paragraph_format.space_before = Pt(before)
        p.paragraph_format.space_after = Pt(after)
        r = p.add_run(text)
        r.bold = bold
        r.font.size = Pt(size)
        r.font.color.rgb = RGBColor(*color)
        return p

    def heading(cell, text: str, before: int = 10):
        para(cell, text.upper(), size=8, bold=True, before=before, after=4)

    para(sidebar, data.get("full_name", ""), size=18, bold=True, after=8)
    if data.get("contact"):
        heading(sidebar, "Contact", 8)
        for part in str(data["contact"]).split(" | "):
            para(sidebar, part, size=8, after=1)

    skills = data.get("skills") or []
    if skills:
        heading(sidebar, "Skills", 12)
        for skill in skills[:14]:
            para(sidebar, skill, size=8, after=1)

    if data.get("certifications"):
        heading(sidebar, "Certifications", 12)
        for cert in data["certifications"]:
            para(sidebar, cert, size=8, after=1)

    para(main, data.get("full_name", ""), size=24, bold=True, after=2)
    if data.get("summary"):
        heading(main, "Professional Summary", 10)
        para(main, data["summary"], size=10, after=6)

    if data.get("experience"):
        heading(main, "Experience", 10)
        for exp in data["experience"]:
            para(main, f"{exp.get('title', '')} - {exp.get('company', '')}", size=10, bold=True, after=0)
            para(main, exp.get("duration", ""), size=8, color=(0x4B, 0x55, 0x63), after=2)
            for bullet in exp.get("achievements", []):
                p = main.add_paragraph(style="List Bullet")
                p.paragraph_format.space_after = Pt(1)
                p.add_run(bullet).font.size = Pt(9)

    if data.get("projects"):
        heading(main, "Projects", 10)
        for proj in data["projects"]:
            para(main, proj.get("name", ""), size=10, bold=True, after=1)
            desc = proj.get("description", "")
            if desc:
                para(main, desc, size=9, after=2)

    if data.get("education"):
        heading(main, "Education", 10)
        for edu in data["education"]:
            para(main, edu.get("degree", ""), size=10, bold=True, after=0)
            para(main, f"{edu.get('institution', '')} | {edu.get('year', '')}", size=9, after=2)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _docx_t1(data: NormalizedData) -> bytes:
    """Classic Professional — blue, centered, standard."""
    return _docx_base(data, name_size=22, name_centered=True, accent_rgb=(0x1D, 0x4E, 0xD8))


def _docx_t2(data: NormalizedData) -> bytes:
    """Compact ATS — teal, slightly smaller fonts, tighter margins."""
    return _docx_base(data, name_size=20, name_centered=True,
                      accent_rgb=(0x11, 0x18, 0x27), compact=True,
                      body_font_size=9, section_font_size=10)


def _docx_t3(data: NormalizedData) -> bytes:
    """Modern ATS Professional — violet, left-aligned name."""
    return _docx_base(data, name_size=24, name_centered=False, accent_rgb=(0x1D, 0x4E, 0xD8))


def _docx_t4(data: NormalizedData) -> bytes:
    """Clean Minimal — dark slate, thin HR, lowercase sections."""
    return _docx_base(data, name_size=20, name_centered=True,
                      accent_rgb=(0x11, 0x18, 0x27), hr_thickness="4",
                      section_uppercase=False, section_font_size=10)


def _docx_t5(data: NormalizedData) -> bytes:
    """Technical Engineer — green, skills-first, grid layout."""
    return _docx_base(data, name_size=22, name_centered=True,
                      accent_rgb=(0x11, 0x18, 0x27), hr_thickness="8",
                      section_font_size=11)


def _docx_t6(data: NormalizedData) -> bytes:
    """Compact One-Page ATS — azure blue, very compact."""
    return _docx_sidebar(data)


def _docx_t7(data: NormalizedData) -> bytes:
    """Executive Professional — dark navy, no bright colors."""
    return _docx_base(data, name_size=22, name_centered=False,
                      accent_rgb=(0x1D, 0x4E, 0xD8), skills_first=True,
                      show_skills_grid=True)


def _docx_t8(data: NormalizedData) -> bytes:
    """Skills-First Hybrid — cyan, skills second after name."""
    return _docx_base(data, name_size=22, name_centered=False,
                      accent_rgb=(0x1E, 0x3A, 0x8A), projects_second=True)


def _docx_t9(data: NormalizedData) -> bytes:
    """Project-Heavy Developer — rose, projects before experience."""
    return _docx_base(data, name_size=24, name_centered=False,
                      accent_rgb=(0x11, 0x18, 0x27), hr_thickness="8",
                      section_font_size=11)


def _docx_t10(data: NormalizedData) -> bytes:
    """Elegant Corporate ATS — amber, left-aligned, thin borders."""
    return _docx_base(data, name_size=20, name_centered=False,
                      accent_rgb=(0x1D, 0x4E, 0xD8), hr_thickness="4",
                      section_uppercase=False, compact=True,
                      body_font_size=9, section_font_size=10)


# ═══════════════════════════════════════════════════════════════════════════════
# PDF BUILDERS
# ═══════════════════════════════════════════════════════════════════════════════

def _pdf_base(
    data: NormalizedData,
    *,
    accent_hex: str = "#1a56db",
    secondary_hex: str = "#64748b",
    divider_hex: str = "#e2e8f0",
    name_size: int = 22,
    name_align: int = 1,       # 1=center, 0=left
    section_size: int = 11,
    body_size: int = 10,
    compact: bool = False,
    skills_first: bool = False,
    projects_second: bool = False,
    section_label_transform: Callable[[str], str] = str.upper,
    hr_before_section: bool = True,
    bold_divider: bool = False,
    left_margin: float = 1.0,
    right_margin: float = 1.0,
) -> bytes:
    """
    Shared PDF builder for all 10 template variants.
    Uses ReportLab's platypus flowable engine.
    """
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable
    from reportlab.lib.units import inch
    from reportlab.lib import colors

    ACCENT = colors.HexColor(accent_hex)
    GRAY = colors.HexColor(secondary_hex)
    LIGHT = colors.HexColor(divider_hex)

    top_m = 0.55 * inch if compact else 0.75 * inch
    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf, pagesize=letter,
        rightMargin=right_margin * inch, leftMargin=left_margin * inch,
        topMargin=top_m, bottomMargin=top_m,
    )

    space_before_section = 6 if compact else 10
    space_after_section = 2 if compact else 3

    # Root-cause fix: without explicit `leading`, ReportLab inherits ~12pt from
    # the base Normal style. For a 22pt name that means the paragraph bounding
    # box is only 12pt tall — the contact Paragraph is placed inside the name
    # glyph, causing visible overlap. Set leading = fontSize + generous cap so
    # the bounding box always encloses the tallest possible glyph.
    name_style = ParagraphStyle(
        "Name",
        fontSize=name_size,
        fontName="Helvetica-Bold",
        leading=int(name_size * 1.4),   # ≥ fontSize; e.g. 22pt → 30pt line box
        alignment=name_align,
        spaceBefore=0,
        spaceAfter=4,                   # clear gap below name before contact
    )
    contact_style = ParagraphStyle(
        "Contact",
        fontSize=body_size - 1,
        leading=int((body_size - 1) * 1.4),
        alignment=name_align,
        textColor=GRAY,
        spaceBefore=0,
        spaceAfter=10 if compact else 12,
    )
    section_style = ParagraphStyle(
        "Section", fontSize=section_size, fontName="Helvetica-Bold",
        textColor=ACCENT, spaceBefore=space_before_section,
        spaceAfter=space_after_section,
    )
    body_style = ParagraphStyle(
        "Body", fontSize=body_size, leading=13 if compact else 14, spaceAfter=3,
    )
    bullet_style = ParagraphStyle(
        "Bullet", fontSize=body_size, leading=13 if compact else 14,
        leftIndent=14, spaceAfter=2,
    )
    job_title_style = ParagraphStyle(
        "JobTitle", fontSize=body_size, fontName="Helvetica-Bold",
        spaceAfter=1, spaceBefore=5 if compact else 6,
    )
    italic_style = ParagraphStyle(
        "Italic", fontSize=body_size - 1, fontName="Helvetica-Oblique",
        textColor=GRAY, spaceAfter=2,
    )
    proj_name_style = ParagraphStyle(
        "ProjName", fontSize=body_size - 1, fontName="Helvetica-BoldOblique",
        textColor=GRAY, spaceBefore=3, spaceAfter=1,
    )

    story: List[Any] = []

    def section(title: str):
        label = section_label_transform(title)
        story.append(Paragraph(label, section_style))
        if hr_before_section:
            thickness = 1.5 if bold_divider else 0.5
            story.append(HRFlowable(width="100%", thickness=thickness, color=LIGHT, spaceAfter=3))

    story.append(Paragraph(data.get("full_name", ""), name_style))
    story.append(Spacer(1, 2))   # belt-and-suspenders: explicit gap between name and contact
    story.append(Paragraph(data.get("contact", ""), contact_style))
    story.append(HRFlowable(width="100%", thickness=1.5, color=ACCENT, spaceBefore=2, spaceAfter=6))

    def _add_summary():
        if data.get("summary"):
            section("Professional Summary")
            story.append(Paragraph(data["summary"], body_style))

    def _add_skills():
        tech = data.get("technical_skills") or []
        prof = data.get("professional_skills") or []
        all_skills = data.get("skills") or tech + prof
        if all_skills:
            section("Skills")
            if tech and prof:
                story.append(Paragraph("<b>Technical:</b> " + " &bull; ".join(tech), body_style))
                story.append(Paragraph("<b>Professional:</b> " + " &bull; ".join(prof), body_style))
            else:
                story.append(Paragraph(" &bull; ".join(all_skills), body_style))

    def _add_experience():
        if data.get("experience"):
            section("Experience")
            for exp in data["experience"]:
                story.append(Paragraph(
                    f"<b>{exp.get('title', '')} — {exp.get('company', '')}</b>",
                    job_title_style,
                ))
                story.append(Paragraph(f"<i>{exp.get('duration', '')}</i>", italic_style))
                sub_projects = exp.get("projects") or []
                if sub_projects:
                    for proj in sub_projects:
                        proj_name = proj.get("name", "")
                        if proj_name:
                            story.append(Paragraph(proj_name, proj_name_style))
                        for bullet in proj.get("bullets") or []:
                            story.append(Paragraph(f"&bull; {bullet}", bullet_style))
                else:
                    for ach in exp.get("achievements", []):
                        story.append(Paragraph(f"&bull; {ach}", bullet_style))
                story.append(Spacer(1, 3 if compact else 4))

    def _add_education():
        if data.get("education"):
            section("Education")
            for edu in data["education"]:
                story.append(Paragraph(f"<b>{edu.get('degree', '')}</b>", job_title_style))
                story.append(Paragraph(
                    f"{edu.get('institution', '')} | {edu.get('year', '')}", body_style
                ))

    def _add_projects():
        if data.get("projects"):
            section("Projects")
            for proj in data["projects"]:
                story.append(Paragraph(f"<b>{proj.get('name', '')}</b>", job_title_style))
                desc = proj.get("description", "")
                if desc:
                    story.append(Paragraph(desc, body_style))
                techs = proj.get("technologies", [])
                if techs:
                    story.append(Paragraph(
                        f"<i>Technologies: {', '.join(techs)}</i>", italic_style
                    ))

    def _add_certifications():
        if data.get("certifications"):
            section("Certifications")
            for cert in data["certifications"]:
                story.append(Paragraph(f"&bull; {cert}", bullet_style))

    # ── Section ordering ──────────────────────────────────────────────────────
    if skills_first:
        _add_summary()
        _add_skills()
        _add_experience()
        _add_projects()
    elif projects_second:
        _add_summary()
        _add_projects()
        _add_experience()
        _add_skills()
    else:
        _add_summary()
        _add_skills()
        _add_experience()

    _add_education()

    if not skills_first and not projects_second:
        _add_projects()

    _add_certifications()

    doc.build(story)
    return buf.getvalue()


# ─── 10 PDF template dispatchers ─────────────────────────────────────────────

def _pdf_sidebar(data: NormalizedData) -> bytes:
    """Two-column PDF layout with contact/skills sidebar and main content."""
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    from reportlab.lib.units import inch
    from reportlab.lib import colors

    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf,
        pagesize=letter,
        rightMargin=0.55 * inch,
        leftMargin=0.55 * inch,
        topMargin=0.55 * inch,
        bottomMargin=0.55 * inch,
    )

    ink = colors.HexColor("#111827")
    blue = colors.HexColor("#1d4ed8")
    soft_blue = colors.HexColor("#dbeafe")
    muted = colors.HexColor("#4b5563")

    sidebar_head = ParagraphStyle("SidebarHead", fontSize=7.5, leading=9, fontName="Helvetica-Bold", textColor=ink, spaceBefore=8, spaceAfter=4)
    sidebar_body = ParagraphStyle("SidebarBody", fontSize=7.5, leading=10, textColor=ink, spaceAfter=2)
    name_style = ParagraphStyle("SidebarName", fontSize=24, leading=30, fontName="Helvetica-Bold", textColor=ink, spaceAfter=4)
    section_style = ParagraphStyle("SidebarSection", fontSize=9, leading=11, fontName="Helvetica-Bold", textColor=blue, spaceBefore=10, spaceAfter=5)
    body_style = ParagraphStyle("SidebarMainBody", fontSize=9.5, leading=13, textColor=ink, spaceAfter=4)
    muted_style = ParagraphStyle("SidebarMuted", fontSize=8, leading=10, textColor=muted, spaceAfter=2)
    bullet_style = ParagraphStyle("SidebarBullet", fontSize=9, leading=12, leftIndent=10, textColor=ink, spaceAfter=2)
    title_style = ParagraphStyle("SidebarTitle", fontSize=10, leading=12, fontName="Helvetica-Bold", textColor=ink, spaceAfter=1, spaceBefore=4)

    left: List[Any] = [Paragraph(f"<b>{data.get('full_name', '')}</b>", ParagraphStyle("LeftName", parent=sidebar_body, fontSize=15, leading=18, fontName="Helvetica-Bold"))]
    if data.get("contact"):
        left.append(Paragraph("CONTACT", sidebar_head))
        for part in str(data["contact"]).split(" | "):
            left.append(Paragraph(part, sidebar_body))

    skills = data.get("skills") or []
    if skills:
        left.append(Paragraph("SKILLS", sidebar_head))
        for skill in skills[:14]:
            left.append(Paragraph(skill, sidebar_body))

    if data.get("certifications"):
        left.append(Paragraph("CERTIFICATIONS", sidebar_head))
        for cert in data["certifications"]:
            left.append(Paragraph(cert, sidebar_body))

    right: List[Any] = [Paragraph(data.get("full_name", ""), name_style)]
    if data.get("summary"):
        right.extend([Paragraph("PROFESSIONAL SUMMARY", section_style), Paragraph(data["summary"], body_style)])

    if data.get("experience"):
        right.append(Paragraph("EXPERIENCE", section_style))
        for exp in data["experience"]:
            right.append(Paragraph(f"{exp.get('title', '')} - {exp.get('company', '')}", title_style))
            right.append(Paragraph(exp.get("duration", ""), muted_style))
            for bullet in exp.get("achievements", []):
                right.append(Paragraph(f"- {bullet}", bullet_style))
            right.append(Spacer(1, 3))

    if data.get("projects"):
        right.append(Paragraph("PROJECTS", section_style))
        for proj in data["projects"]:
            right.append(Paragraph(proj.get("name", ""), title_style))
            desc = proj.get("description", "")
            if desc:
                right.append(Paragraph(desc, body_style))

    if data.get("education"):
        right.append(Paragraph("EDUCATION", section_style))
        for edu in data["education"]:
            right.append(Paragraph(edu.get("degree", ""), title_style))
            right.append(Paragraph(f"{edu.get('institution', '')} | {edu.get('year', '')}", muted_style))

    table = Table([[left, right]], colWidths=[1.85 * inch, 4.95 * inch])
    table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (0, 0), soft_blue),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("LEFTPADDING", (0, 0), (0, 0), 14),
        ("RIGHTPADDING", (0, 0), (0, 0), 12),
        ("TOPPADDING", (0, 0), (0, 0), 14),
        ("BOTTOMPADDING", (0, 0), (0, 0), 14),
        ("LEFTPADDING", (1, 0), (1, 0), 18),
        ("RIGHTPADDING", (1, 0), (1, 0), 4),
        ("TOPPADDING", (1, 0), (1, 0), 4),
        ("LINEBEFORE", (1, 0), (1, 0), 1.5, blue),
    ]))
    doc.build([table])
    return buf.getvalue()


def _pdf_t1(data: NormalizedData) -> bytes:
    """Classic Professional — blue, centered."""
    return _pdf_base(data, accent_hex="#1d4ed8", name_align=1)


def _pdf_t2(data: NormalizedData) -> bytes:
    """Compact ATS — teal, tighter layout."""
    return _pdf_base(data, accent_hex="#111827", secondary_hex="#475569",
                     name_size=20, body_size=9, section_size=10,
                     compact=True, name_align=1)


def _pdf_t3(data: NormalizedData) -> bytes:
    """Modern ATS Professional — violet, left-aligned name."""
    return _pdf_base(data, accent_hex="#1d4ed8", divider_hex="#dbeafe",
                     name_size=24, name_align=0, bold_divider=True)


def _pdf_t4(data: NormalizedData) -> bytes:
    """Clean Minimal — dark slate, thin dividers, title-case labels."""
    return _pdf_base(
        data, accent_hex="#111827", secondary_hex="#6b7280",
        divider_hex="#f1f5f9", name_size=20, section_size=10,
        section_label_transform=str.title,
    )


def _pdf_t5(data: NormalizedData) -> bytes:
    """Technical Engineer — green, skills lead."""
    return _pdf_base(data, accent_hex="#111827", secondary_hex="#475569",
                     divider_hex="#cbd5e1", name_size=22, bold_divider=True,
                     name_align=1, section_size=11)


def _pdf_t6(data: NormalizedData) -> bytes:
    """Compact One-Page ATS — azure blue, very dense."""
    return _pdf_sidebar(data)


def _pdf_t7(data: NormalizedData) -> bytes:
    """Executive Professional — near-black, thick dividers."""
    return _pdf_base(data, accent_hex="#1d4ed8", divider_hex="#dbeafe",
                     name_align=0, skills_first=True)


def _pdf_t8(data: NormalizedData) -> bytes:
    """Skills-First Hybrid — cyan, left-aligned, skills lead."""
    return _pdf_base(data, accent_hex="#1e3a8a", divider_hex="#dbeafe",
                     name_align=1, projects_second=True)


def _pdf_t9(data: NormalizedData) -> bytes:
    """Project-Heavy Developer — rose, projects before experience."""
    return _pdf_base(data, accent_hex="#111827", secondary_hex="#475569",
                     divider_hex="#cbd5e1", name_size=24, bold_divider=True,
                     name_align=0, section_size=11)


def _pdf_t10(data: NormalizedData) -> bytes:
    """Elegant Corporate ATS — amber, left-aligned, thin dividers."""
    return _pdf_base(data, accent_hex="#1d4ed8", secondary_hex="#475569",
                     divider_hex="#dbeafe", name_size=20, name_align=0,
                     section_label_transform=str.title, section_size=10,
                     compact=True, body_size=9)


# ═══════════════════════════════════════════════════════════════════════════════
# DISPATCH TABLES
# ═══════════════════════════════════════════════════════════════════════════════

_PDF_BUILDERS: Dict[str, Callable[[NormalizedData], bytes]] = {
    "template_1": _pdf_t1,
    "template_2": _pdf_t2,
    "template_3": _pdf_t3,
    "template_4": _pdf_t4,
    "template_5": _pdf_t5,
    "template_6": _pdf_t6,
    "template_7": _pdf_t7,
    "template_8": _pdf_t8,
    "template_9": _pdf_t9,
    "template_10": _pdf_t10,
}

_DOCX_BUILDERS: Dict[str, Callable[[NormalizedData], bytes]] = {
    "template_1": _docx_t1,
    "template_2": _docx_t2,
    "template_3": _docx_t3,
    "template_4": _docx_t4,
    "template_5": _docx_t5,
    "template_6": _docx_t6,
    "template_7": _docx_t7,
    "template_8": _docx_t8,
    "template_9": _docx_t9,
    "template_10": _docx_t10,
}

# New slug IDs → legacy template_N keys
_SLUG_ALIASES: Dict[str, str] = {
    "classic-professional":    "template_1",
    "compact-ats":             "template_2",
    "modern-ats-professional": "template_3",
    "minimal-one-column":      "template_4",
    "executive-clean":         "template_5",
    "sidebar-professional":    "template_6",
    "technical-skills-first":  "template_7",
    "project-portfolio":       "template_8",
    "senior-leadership":       "template_9",
    "corporate-compact":       "template_10",
    # Backward-compatible aliases for previously shipped slugs.
    "clean-minimal":           "template_4",
    "technical-engineer":      "template_7",
    "compact-one-page":        "template_10",
    "executive-professional":  "template_5",
    "skills-first-hybrid":     "template_7",
    "project-heavy-developer": "template_8",
    "elegant-corporate":       "template_10",
}


def build_pdf(template_id: str, data: NormalizedData) -> bytes:
    """
    Build a PDF for the given template_id.
    Accepts both legacy 'template_N' IDs and new slug IDs.
    Falls back to template_1 (Classic Professional) if template_id is unknown.
    """
    resolved = _SLUG_ALIASES.get(template_id, template_id)
    builder = _PDF_BUILDERS.get(resolved, _pdf_t1)
    return builder(data)


def build_docx(template_id: str, data: NormalizedData) -> bytes:
    """
    Build a DOCX for the given template_id.
    Accepts both legacy 'template_N' IDs and new slug IDs.
    Falls back to template_1 (Classic Professional) if template_id is unknown.
    """
    resolved = _SLUG_ALIASES.get(template_id, template_id)
    builder = _DOCX_BUILDERS.get(resolved, _docx_t1)
    return builder(data)
